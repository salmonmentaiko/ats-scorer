"""
Unit tests for core logic in app.py.

Functions are copied directly here to avoid triggering Streamlit's
module-level side-effects (st.set_page_config, st.markdown, Groq(), etc.)
that run on import.
"""

import re
import io
import pytest
from unittest.mock import MagicMock, patch, PropertyMock


# ── Copied function logic (no Streamlit dependency) ──────────────────────────

def extract_keywords(text):
    stopwords = set([
        "and","or","the","a","an","in","on","at","to","for","of","with",
        "is","are","was","were","be","been","have","has","had","will","would",
        "can","could","should","may","might","do","does","did","not","this",
        "that","we","you","i","they","it","as","by","from","our","your",
        "their","all","any","both","each","more","other","such","than","then",
        "so","no","nor","too","very","just","but","if"
    ])
    words = re.findall(r'\b[a-zA-Z][a-zA-Z+#\.]{1,}\b', text.lower())
    return set(w for w in words if w not in stopwords and len(w) > 2)


def score_resume(resume_text, job_text):
    resume_kw = extract_keywords(resume_text)
    job_kw    = extract_keywords(job_text)
    matched   = resume_kw & job_kw
    missing   = job_kw - resume_kw
    score     = round(len(matched) / len(job_kw) * 100) if job_kw else 0
    return score, matched, sorted(missing, key=len, reverse=True)[:15]


def extract_text_from_file(uploaded_file):
    import pdfplumber
    import docx

    if uploaded_file.type == "application/pdf":
        with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif uploaded_file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]:
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    elif uploaded_file.type == "text/plain":
        return uploaded_file.read().decode("utf-8")
    return ""


# ── Tests: extract_keywords ──────────────────────────────────────────────────

class TestExtractKeywords:

    def test_returns_set(self):
        result = extract_keywords("Python developer with experience")
        assert isinstance(result, set)

    def test_basic_keywords_extracted(self):
        result = extract_keywords("Python developer machine learning")
        assert "python" in result
        assert "developer" in result
        assert "machine" in result
        assert "learning" in result

    def test_stopwords_removed(self):
        result = extract_keywords("and or the a an in on at to for of with is are")
        assert result == set(), f"Expected empty set, got {result}"

    def test_short_words_filtered(self):
        # Words of length <= 2 (after the regex requires len > 2) should be absent.
        # The regex already requires at least 2 chars after the first, so single/two-char
        # words (e.g. "go", "be") won't appear.
        result = extract_keywords("go be ok")
        # "be" is a stopword; "go" and "ok" are 2 chars and won't match
        # (regex needs [a-zA-Z][a-zA-Z+#.]{1,} = minimum 2 chars, but len > 2 filter
        # removes 2-char words too)
        assert "go" not in result
        assert "ok" not in result

    def test_empty_string_returns_empty_set(self):
        result = extract_keywords("")
        assert result == set()

    def test_special_characters_ignored(self):
        # Pure numbers should not be captured (regex requires first char alpha)
        result = extract_keywords("123 456 789")
        assert "123" not in result
        assert "456" not in result

    def test_standalone_symbols_ignored(self):
        # Standalone symbols with no leading alpha char should not match
        result = extract_keywords("!!! ??? ### $$$")
        # None of these start with [a-zA-Z] so the regex won't capture them
        assert len(result) == 0

    def test_mixed_case_lowercased(self):
        result = extract_keywords("Python PYTHON python")
        # All should collapse to the same lowercased token
        assert "python" in result
        assert "Python" not in result
        assert "PYTHON" not in result

    def test_tech_keywords_with_plus_chars(self):
        # Known limitation: c++ is not captured because \b splits at '+'.
        # The surrounding word "experience" is still extracted correctly.
        result = extract_keywords("experience with c++")
        assert "experience" in result

    def test_duplicates_collapsed_to_set(self):
        result = extract_keywords("python python python developer developer")
        assert result == {"python", "developer"}

    def test_only_stopwords_returns_empty(self):
        result = extract_keywords("and the or but if so")
        assert result == set()


# ── Tests: score_resume ──────────────────────────────────────────────────────

class TestScoreResume:

    def test_returns_three_tuple(self):
        result = score_resume("python developer", "python engineer")
        assert len(result) == 3

    def test_score_is_int(self):
        score, _, _ = score_resume("python developer", "python engineer")
        assert isinstance(score, int)

    def test_matched_is_set(self):
        _, matched, _ = score_resume("python developer", "python engineer")
        assert isinstance(matched, set)

    def test_missing_is_list(self):
        _, _, missing = score_resume("python developer", "python engineer")
        assert isinstance(missing, list)

    def test_perfect_match_score_100(self):
        text = "python machine learning developer"
        score, matched, missing = score_resume(text, text)
        assert score == 100
        assert missing == []

    def test_no_match_score_0(self):
        score, matched, missing = score_resume(
            "java backend spring boot",
            "python machine learning tensorflow"
        )
        assert score == 0
        assert len(matched) == 0

    def test_partial_match_score(self):
        # Resume has 2 of 4 job keywords
        score, matched, missing = score_resume(
            "python developer",
            "python developer machine learning"
        )
        # job keywords: {python, developer, machine, learning} = 4
        # matched: {python, developer} = 2 → 50%
        assert score == 50
        assert "python" in matched
        assert "developer" in matched

    def test_empty_resume_score_0(self):
        score, matched, missing = score_resume("", "python developer machine learning")
        assert score == 0
        assert matched == set()

    def test_empty_job_description_score_0(self):
        score, matched, missing = score_resume("python developer", "")
        assert score == 0

    def test_both_empty_score_0(self):
        score, matched, missing = score_resume("", "")
        assert score == 0

    def test_missing_capped_at_15(self):
        # Build a job description with more than 15 unique meaningful keywords
        job_words = [
            "accountancy", "analytics", "architecture", "budgeting",
            "collaboration", "compliance", "deployment", "documentation",
            "engineering", "forecasting", "governance", "infrastructure",
            "integration", "kubernetes", "leadership", "mentoring",
            "networking", "optimization", "postgresql", "reporting",
            "scalability", "stakeholders", "statistics", "terraform",
        ]
        job_text = " ".join(job_words)
        score, matched, missing = score_resume("", job_text)
        assert len(missing) <= 15

    def test_missing_sorted_by_length_descending(self):
        # Ensure missing list is sorted longest-first
        score, matched, missing = score_resume(
            "",
            "collaboration infrastructure deployment analytics"
        )
        if len(missing) > 1:
            for i in range(len(missing) - 1):
                assert len(missing[i]) >= len(missing[i + 1])

    def test_score_is_rounded_integer(self):
        # 1 match out of 3 job keywords → 33.33... → rounds to 33
        score, _, _ = score_resume(
            "python",
            "python developer machine"
        )
        assert score == 33

    def test_matched_contains_only_common_keywords(self):
        score, matched, missing = score_resume(
            "python sql javascript",
            "python java sql"
        )
        assert "python" in matched
        assert "sql" in matched
        # javascript is only in resume, not job — should not be in matched
        assert "javascript" not in matched
        # java is only in job, not resume — should not be in matched
        assert "java" not in matched

    def test_missing_does_not_contain_matched(self):
        _, matched, missing = score_resume(
            "python developer analytics",
            "python developer machine learning analytics"
        )
        for kw in matched:
            assert kw not in missing


# ── Tests: extract_text_from_file ────────────────────────────────────────────

class TestExtractTextFromFile:

    def _make_mock_file(self, file_type, content_bytes=b"hello world"):
        mock_file = MagicMock()
        mock_file.type = file_type
        mock_file.read.return_value = content_bytes
        return mock_file

    def test_txt_file_returns_decoded_text(self):
        mock_file = self._make_mock_file("text/plain", b"Hello from TXT file")
        result = extract_text_from_file(mock_file)
        assert result == "Hello from TXT file"

    def test_txt_file_utf8_decoding(self):
        text = "Resume with unicode: café résumé"
        mock_file = self._make_mock_file("text/plain", text.encode("utf-8"))
        result = extract_text_from_file(mock_file)
        assert result == text

    def test_unknown_file_type_returns_empty_string(self):
        mock_file = self._make_mock_file("image/png", b"\x89PNG")
        result = extract_text_from_file(mock_file)
        assert result == ""

    def test_pdf_file_extracts_text(self):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Python Developer Resume"

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        mock_file = self._make_mock_file("application/pdf", b"%PDF-fake")

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_file(mock_file)

        assert result == "Python Developer Resume"

    def test_pdf_multiple_pages_joined_with_newline(self):
        page1 = MagicMock()
        page1.extract_text.return_value = "Page one content"
        page2 = MagicMock()
        page2.extract_text.return_value = "Page two content"

        mock_pdf = MagicMock()
        mock_pdf.pages = [page1, page2]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        mock_file = self._make_mock_file("application/pdf", b"%PDF-fake")

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_file(mock_file)

        assert result == "Page one content\nPage two content"

    def test_pdf_page_with_none_text_treated_as_empty(self):
        page = MagicMock()
        page.extract_text.return_value = None

        mock_pdf = MagicMock()
        mock_pdf.pages = [page]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        mock_file = self._make_mock_file("application/pdf", b"%PDF-fake")

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = extract_text_from_file(mock_file)

        assert result == ""

    def test_docx_file_extracts_paragraphs(self):
        para1 = MagicMock()
        para1.text = "Experienced software engineer"
        para2 = MagicMock()
        para2.text = "Proficient in Python and SQL"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [para1, para2]

        mock_file = self._make_mock_file(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            b"PK\x03\x04fake-docx"
        )

        with patch("docx.Document", return_value=mock_doc):
            result = extract_text_from_file(mock_file)

        assert result == "Experienced software engineer\nProficient in Python and SQL"

    def test_doc_legacy_format_extracts_paragraphs(self):
        para = MagicMock()
        para.text = "Legacy Word document"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [para]

        mock_file = self._make_mock_file("application/msword", b"\xD0\xCF\x11fake")

        with patch("docx.Document", return_value=mock_doc):
            result = extract_text_from_file(mock_file)

        assert result == "Legacy Word document"

    def test_txt_empty_file_returns_empty_string(self):
        mock_file = self._make_mock_file("text/plain", b"")
        result = extract_text_from_file(mock_file)
        assert result == ""
